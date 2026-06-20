"""Rebuild the canonical Supplementary Text S4 (Nested-CV Pipeline) docx in
GenomeMedicine_TNT_v2/supplementary_texts/, replacing the stale 2026-04-29
copy that still claims '0.65-0.69 reference value used throughout the
manuscript'.

The new S4 makes 28-feature reference set / ElasticNet AUC = 0.745 the single
reference, in line with Manuscript_v4.docx (§3.5 / Fig 4E / Supp Fig S5A /
Abstract). Adds the 5-scenario cherry-picking table (data from
260418_add/nested_cv_drop_vs_swap.tsv) and a regenerated per-fold feature
stability table on the 28-feature set (data from
260418_add/nested_cv_28feat_feature_selection_frequency.tsv).

Cross-references match canonical v2 numbering: §3.5, Fig 4E, Supp Fig S5A,
Supp Text S3 (CD8 audit).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

OUT = Path("/data/data/TNT/analysis/GenomeMedicine_TNT_v3/supplementary_texts/New_Supp_Text_S4_nested_CV_pipeline.docx")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
for section in doc.sections:
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = 914400

def H(text, level=2):
    h = doc.add_heading(level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
    h.add_run(text)

def P(*runs):
    """Add a paragraph from a list of (text, style) tuples where style ∈ {'', 'b', 'i', 'c'}."""
    p = doc.add_paragraph()
    for text, kind in runs:
        run = p.add_run(text)
        if kind == 'b':
            run.bold = True
        elif kind == 'i':
            run.italic = True
        elif kind == 'c':
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
    return p

def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)

def TABLE(header, rows, bold_rows=()):
    tbl = doc.add_table(rows=1+len(rows), cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for j, h in enumerate(header):
        cp = tbl.rows[0].cells[j].paragraphs[0]
        run = cp.add_run(h)
        run.bold = True
    for ri, row in enumerate(rows):
        for j, c in enumerate(row[:len(header)]):
            cp = tbl.rows[ri+1].cells[j].paragraphs[0]
            run = cp.add_run(c)
            if ri in bold_rows:
                run.bold = True
    doc.add_paragraph()

# ------------- Title + header -------------
H("Supplementary Text S4 — Nested-CV Pipeline on the 28-Feature Reference Set", level=1)
P(("This document details the leakage-free cross-validation procedure underlying the headline classifier AUC reported throughout the manuscript (Abstract, §3.5, Fig 4E, Supp Fig S5A): ", ''),
  ("ElasticNet outer-LOOCV AUC = 0.745", 'b'),
  (" (95 % bootstrap CI 0.56–0.89) on the ", ''),
  ("28-feature reference set", 'b'),
  (" (the 37-column integrated master table — 29 numeric model inputs after dropping 8 ID/categorical columns — minus the contaminated ", ''),
  ("CD8_proliferation", 'c'),
  (" feature; see S4.1 below and Supp Text S3 for the contamination diagnosis). Four additional pre-specified scenarios are reported in S4.4 as a cherry-picking sanity check.", ''))

# ------------- S4.1 Motivation -------------
H("S4.1 Motivation and reference scenario", level=2)
P(("An earlier internal analysis selected features on the full training set and then evaluated a LASSO classifier by leave-one-out cross-validation, yielding AUC = 0.755. Because the feature-selection step had access to the held-out fold, this estimate is optimistically biased and is not used as a performance claim anywhere in the manuscript; it is reported here only for transparency.", ''))
P(("We re-implemented the predictor with full nesting (outer LOOCV, inner 5-fold hyperparameter tuning, feature pre-selection refit inside each inner training fold) to obtain a leakage-free estimate. The original 37-column master table (29 numeric model inputs after dropping 8 ID/categorical columns; Table S1) contained a ", ''),
  ("CD8_proliferation", 'c'),
  (" feature whose underlying gene panel was independently shown (Supp Text S3) to be contaminated by cell-cycle genes — a contamination that mis-attributed cell-cycle/proliferation signal to the CD8 axis and conflated the two largest tumor-intrinsic and immune blocks of the predictor. The contamination diagnosis was made on biological grounds (gene-panel composition audit), not on AUC, and is documented in S3.", ''))
P(("Once this single contaminated feature is removed, the cleaned ", ''),
  ("28-feature reference set", 'b'),
  (" is the analytically correct feature set for the predictor, and its honest nested-CV ElasticNet AUC = ", ''),
  ("0.745", 'b'),
  (" (95 % bootstrap CI 0.56–0.89) is ", ''),
  ("the reference value used throughout the manuscript", 'b'),
  (" (Abstract, §3.5, Fig 4E, Supp Fig S5A). The corresponding LASSO AUC on the same 28-feature set is 0.716 (95 % CI 0.52–0.88) and is reported jointly. Earlier internal numbers from the contaminated 37-column master baseline — i.e. the same 29 model inputs ", ''),
  ("with", 'i'),
  (" ", ''),
  ("CD8_proliferation", 'c'),
  (" retained — (LASSO 0.650 / ElasticNet 0.686) appear only in the cherry-picking sanity-check table in S4.4 as the historical pre-S3 baseline; they are ", ''),
  ("not", 'i'),
  (" a reference value and are not used as a performance claim in the main text.", ''))

# ------------- S4.2 Pseudocode -------------
H("S4.2 Pseudocode", level=2)
P(("The pseudocode below is illustrated for LASSO; ElasticNet (the headline learner) and RandomForest (a comparator, S4.5) use the same nesting structure with their respective hyperparameter grids.", ''))
CODE("""Input:
  X  : 35 × 28 feature matrix (clinical + WES + RNA features on the
       28-feature reference set, Table S1)
  y  : 35-vector binary response (good = 1, bad = 0)
  K_outer = leave-one-out (n_outer = 35)
  K_inner = 5-fold stratified
  k_grid = {5, 8, 12}       # number of features to retain
  C_grid = {0.1, 0.3, 1, 3} # inverse L2 strength

For i in 1..n_outer:
  X_train, X_test = X[-i], X[i]
  y_train, y_test = y[-i], y[i]

  best_score, best_k, best_C = -inf, NA, NA
  For k in k_grid:
    For C in C_grid:
      inner_scores = []
      For each inner fold (X_inner_tr, y_inner_tr, X_inner_va, y_inner_va):
        # FEATURE SELECTION inside inner training set only
        feats = SelectKBest(f_classif, k=k).fit(X_inner_tr, y_inner_tr)
        X_inner_tr_sel = feats.transform(X_inner_tr)
        X_inner_va_sel = feats.transform(X_inner_va)
        clf = LogisticRegression(penalty='l1', C=C, solver='liblinear')
        clf.fit(X_inner_tr_sel, y_inner_tr)
        inner_scores.append(roc_auc_score(y_inner_va,
                            clf.predict_proba(X_inner_va_sel)[:,1]))
      mean_inner = mean(inner_scores)
      If mean_inner > best_score:
        best_score, best_k, best_C = mean_inner, k, C

  # Refit on full outer training set with selected hyperparameters
  feats = SelectKBest(f_classif, k=best_k).fit(X_train, y_train)
  X_train_sel = feats.transform(X_train)
  X_test_sel  = feats.transform(X_test)
  clf = LogisticRegression(penalty='l1', C=best_C, solver='liblinear')
  clf.fit(X_train_sel, y_train)
  outer_pred[i] = clf.predict_proba(X_test_sel)[:,1]
  outer_features[i] = list of selected features

Output (28-feature reference set, ElasticNet, l1_ratio grid):
  outer_auc = roc_auc_score(y, outer_pred)            # = 0.745
  outer_auc_CI = bootstrap_2000(outer_auc) = [0.56, 0.89]
  feature_stability = frequency of each feature across n_outer folds""")

# ------------- S4.3 Feature stability -------------
H("S4.3 Feature stability across outer folds (28-feature reference set)", level=2)
P(("Across the 35 outer LOOCV folds of the 28-feature reference set, the features most frequently retained by the inner-tuned ", ''),
  ("SelectKBest", 'c'),
  (" step were a tight DNA-damage / cell-cycle / DNA-repair core (selection counts under the ElasticNet learner; LASSO counts in the right-most column):", ''))
TABLE(
    header=["Feature", "ElasticNet", "LASSO"],
    rows=[
        ["DNA Double-Strand Break Repair (R-HSA-5693532)", "35 / 35", "35 / 35"],
        ["HDR via Homologous Recombination (R-HSA-5685942)", "35 / 35", "35 / 35"],
        ["DNA Repair (R-HSA-73894)", "35 / 35", "35 / 35"],
        ["G2-M Checkpoint", "35 / 35", "35 / 35"],
        ["E2F Targets", "34 / 35", "35 / 35"],
        ["MYC Targets V2", "27 / 35", "31 / 35"],
        ["frac_amp (genomic amplification fraction)", "27 / 35", "30 / 35"],
        ["TGF-β Signaling", "27 / 35", "31 / 35"],
        ["Stemness mRNAsi proxy", "26 / 35", "31 / 35"],
        ["MHC II", "25 / 35", "31 / 35"],
        ["MSI_pct", "19 / 35", "24 / 35"],
        ["Epithelial-Mesenchymal Transition (Hallmark)", "15 / 35", "18 / 35"],
        ["SBS5", "14 / 35", "19 / 35"],
    ],
    bold_rows=range(5),
)
P(("A tight five-feature DNA-damage-and-cell-cycle core (DSB repair, HDR, generic DNA repair, G2-M checkpoint, E2F targets) is selected in ", ''),
  ("34–35 of the 35 outer folds", 'b'),
  (" under both learners; this five-feature core is the analytical basis of the tumor-intrinsic predictor described in §3.5 of the main text. MYC Targets V2, ", ''),
  ("frac_amp", 'c'),
  (", TGF-β Signaling, Stemness mRNAsi, and MHC II are stable secondary contributors (25–31 / 35 under at least one learner). MSI_pct, EMT Hallmark, and SBS5 are intermittent (14–24 / 35). The remaining 15 features in the 28-feature reference set (", ''),
  ("age", 'c'), (", ", ''),
  ("TMB_nonsyn_per_Mb", 'c'), (", ", ''),
  ("n_nonsyn", 'c'), (", ", ''),
  ("CIN", 'c'), (", ", ''),
  ("frac_del", 'c'), (", ", ''),
  ("MMR_prop", 'c'), (", ", ''),
  ("CD8_activation", 'c'), (", ", ''),
  ("Antigen_presentation", 'c'), (", ", ''),
  ("NLRC5_HLA_IFNG", 'c'), (", ", ''),
  ("TLS_Cabrita", 'c'), (", ", ''),
  ("IFNg_Ayers_18", 'c'), (", ", ''),
  ("TGFb_Mariathasan", 'c'), (", ", ''),
  ("EMT_Mak", 'c'), (", Buffa Hypoxia, Hallmark Hypoxia) are selected in 0–2 / 35 folds and contribute negligibly to the predictor. Per-fold selected feature lists and full counts are tabulated in ", ''),
  ("260418_add/nested_cv_28feat_features_per_fold.tsv", 'c'),
  (" and ", ''),
  ("260418_add/nested_cv_28feat_feature_selection_frequency.tsv", 'c'), (".", ''))

# ------------- S4.4 Cherry-picking sanity check -------------
H("S4.4 Cherry-picking sanity check across five pre-specified feature sets", level=2)
P(("The 28-feature reference set (S4.1) was selected on biological grounds — removal of the ", ''),
  ("CD8_proliferation", 'c'),
  (" feature once cell-cycle contamination was diagnosed in S3 — ", ''),
  ("not", 'i'),
  (" on AUC. To rule out the alternative explanation that the AUC gain is a feature-set cherry-picking artifact, we evaluated five pre-specified scenarios under an identical nested outer-LOOCV / inner 5-fold pipeline (data: ", ''),
  ("260418_add/nested_cv_drop_vs_swap.tsv", 'c'), ("):", ''))
TABLE(
    header=["Scenario", "Description", "LASSO AUC [95 % CI]", "ElasticNet AUC [95 % CI]"],
    rows=[
        ["baseline_37 (pre-S3 baseline)", "original 37-column master, 29 model inputs, retains contaminated CD8_proliferation", "0.650 [0.45, 0.83]", "0.686 [0.50, 0.86]"],
        ["drop_cd8prolif_36 (reference)", "baseline minus contaminated CD8_proliferation — 28-feature reference set", "0.716 [0.52, 0.88]", "0.745 [0.56, 0.89]"],
        ["swap_cd8_37", "replace CD8_proliferation with a curated CD8-cytotoxic score", "0.716 [0.52, 0.88]", "0.745 [0.56, 0.89]"],
        ["drop_prolif_add_3_39", "drop CD8_proliferation, add 3 cytotoxic / IFN-γ / TLS scores", "0.716 [0.52, 0.88]", "0.735 [0.55, 0.89]"],
        ["add_immune_40", "retain CD8_proliferation, add 3 ssGSEA immune signatures", "0.650 [0.45, 0.83]", "0.686 [0.50, 0.86]"],
    ],
    bold_rows=[1],
)
P(("Two observations support the biological-rather-than-empirical selection of the 28-feature reference:", ''))
P(("1. The two scenarios that explicitly remove the contaminated feature (", ''),
  ("drop_cd8prolif_36", 'c'), (", ", ''),
  ("swap_cd8_37", 'c'),
  (") converge to the same maximum (ElasticNet 0.745). Replacing the contaminated feature with a clean CD8-cytotoxic score gives the same AUC, indicating the gain is from removing the contamination, not from any specific replacement.", ''))
P(("2. Adding clean immune features ", ''),
  ("on top of", 'i'),
  (" the contaminated baseline (", ''),
  ("add_immune_40", 'c'),
  (") does not raise AUC at all (0.650 / 0.686, identical to baseline) — a signal cannot be rescued by addition while contamination remains; only removal works. Adding the same clean features on top of an already-cleaned set (", ''),
  ("drop_prolif_add_3_39", 'c'),
  (") also does not raise AUC further (0.735 vs 0.745), so the headline is not driven by ", ''),
  ("which", 'i'),
  (" additional features were chosen.", ''))
P(("The 28-feature reference is therefore the smallest cleanly-justifiable feature set, and the 0.745 AUC is internally consistent across the two clean scenarios. The pre-S3 baseline (0.650 / 0.686) is included in the table only to make this contrast explicit; it is ", ''),
  ("not", 'i'),
  (" used as a performance claim in the main text.", ''))

# ------------- S4.5 Comparator models -------------
H("S4.5 Comparator models on the 28-feature reference set", level=2)
P(("LASSO outer AUC = 0.716 (95 % CI 0.52–0.88). ElasticNet (", ''),
  ("l1_ratio", 'c'),
  (" ∈ {0.3, 0.5, 0.7, 0.9}; same nested protocol) outer AUC = 0.745 (95 % CI 0.56–0.89) — the headline number. RandomForest (", ''),
  ("n_estimators", 'c'),
  (" = 500, default depth, same nested protocol) gave outer AUC = 0.581 on the 32-input pre-S3 baseline and similar small-N behaviour on the 28-feature reference; SVM-RBF was not tested. LASSO and ElasticNet show overlapping CIs and are reported jointly; RandomForest's lower performance suggests the signal is approximately linear in the selected feature space.", ''))

# ------------- S4.6 Permutation null -------------
H("S4.6 Permutation null", level=2)
P(("A 1,000-permutation null was attempted in which ", ''),
  ("y", 'c'),
  (" labels were shuffled inside the outer loop and the nested pipeline re-run. Wall-clock was prohibitive (estimated 14 hours on the available compute) and the 95 % bootstrap CI of the unpermuted AUC on the 28-feature reference set ([0.56, 0.89]) already approaches 0.5 at the lower bound, transparently bounding the predictor's discriminative power at the small-N discovery stage. Per the project's pragmatism principle, the permutation test was deferred to an external multi-cohort validation effort, the reproducibility of which is the actual basis on which the discovery-stage tumor-intrinsic and immune-axis claims are later defended (§3.10–§3.11).", ''))

# ------------- S4.7 Code -------------
H("S4.7 Code and data", level=2)
P(("scripts/28_nested_cv_permutation.py", 'c'),
  (" — original LASSO / ElasticNet / RandomForest nested pipeline on the contaminated baseline.", ''))
P(("260418_add/05_drop_vs_swap.py", 'c'),
  (" — five-scenario sweep, including the 28-feature reference scenario.", ''))
P(("scripts/54_28feat_feature_stability.py", 'c'),
  (" — per-fold feature-selection frequency on the 28-feature reference set; reproduces the S4.3 table.", ''))
P(("10_ml_predictor/nested_cv_results.tsv", 'c'),
  (", ", ''),
  ("260418_add/nested_cv_drop_vs_swap.tsv", 'c'),
  (", ", ''),
  ("260418_add/probs_drop_cd8prolif_36_{LASSO,ElasticNet}.tsv", 'c'),
  (" — per-scenario / per-subject predictions.", ''))
P(("260418_add/nested_cv_28feat_features_per_fold.tsv", 'c'),
  (", ", ''),
  ("260418_add/nested_cv_28feat_feature_selection_frequency.tsv", 'c'),
  (" — per-fold selected feature lists / 35-fold counts.", ''))
P(("figures/panels_v3/Fig5B_nested_ROC.{pdf,png}", 'c'),
  (" — rendered ROC (legacy filename retained for back-compatibility; corresponds to manuscript Fig 4E ROC).", ''))

doc.save(OUT)
print(f"Wrote {OUT} ({OUT.stat().st_size} bytes)")
